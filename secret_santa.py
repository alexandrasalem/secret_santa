from docx import Document
import random

names = ['Alex', 'Grace', 'Charlie', 'Karen', 'Eric', 'Joanne', 'Delfina', 'Papa and Judy']
#random.shuffle(names)


names_1 = names[:]
names_2 = names[:]
names_3 = names[:]

def accept_permutation(orig, inp):
	condition = True
	for i in range(0, 8):
		if inp[i] ==  orig[i]:
			condition = False
	if condition == True:
		#print(inp)
		return inp

test = None
while test == None:
	random.shuffle(names_1)
	test = accept_permutation(names, names_1)
	#print(names_1)

test2 = None
test3 = None
while test2 == None or test3 == None:
	random.shuffle(names_2)
	test2 = accept_permutation(names, names_2)
	test3 = accept_permutation(test, names_2)
	#print(names_2)

test4 = None
test5 = None
test6 = None
while test4 == None or test5 == None or test6 == None:
	random.shuffle(names_3)
	test4 = accept_permutation(names, names_3)
	test5 = accept_permutation(test, names_3)
	test6 = accept_permutation(test2, names_3)
	test7 = accept_permutation(test3, names_3)
	#print(names_3)

print(names)
print(names_1)
print(names_2)
print(names_3)

for name in names:
    document = Document()
    document.add_heading('Secret Santa', 0)
    document.add_paragraph('You are: ' + name)
    document.add_paragraph('Your assigned people to get presents for: ')
    document.add_paragraph(names_1[names.index(name)], style='List Bullet')
    document.add_paragraph(names_2[names.index(name)], style='List Bullet')
    document.add_paragraph(names_3[names.index(name)], style='List Bullet')
    document.add_paragraph('Yay! Merry Xmas! :)')
    document.save(name +"_secret_santa.docx")

'''
alex_count = 0
grace_count = 0
charlie_count = 0
karen_count = 0
eric_count = 0
joanne_count = 0
delfina_count = 0
papajudy_count = 0

for name in names:
    names_copy = names[:]
    names_copy.remove(name)
    person1_choices = []
    while len(person1_choices) !=3:
        choice = random.choice(names_copy)
        if choice not in person1_choices:
            if choice == 'Alex':
                if alex_count <3:
                    person1_choices.append(choice)
                    alex_count+=1
            elif choice == 'Grace':
                if grace_count <3:
                    person1_choices.append(choice)
                    grace_count+=1
            elif choice == 'Charlie':
                if charlie_count <3:
                    person1_choices.append(choice)
                    charlie_count+=1
            elif choice == 'Karen':
                if karen_count <3:
                    person1_choices.append(choice)
                    karen_count+=1
            elif choice == 'Eric':
                if eric_count <3:
                    person1_choices.append(choice)
                    eric_count+=1
            elif choice == 'Joanne':
                if joanne_count <3:
                    person1_choices.append(choice)
                    joanne_count+=1
            elif choice == 'Delfina':
                if delfina_count <3:
                    person1_choices.append(choice)
                    delfina_count+=1
            else:
                if papajudy_count <3:
                    person1_choices.append(choice)
                    papajudy_count+=1
        print(person1_choices)
    print(name)
    print(person1_choices)

    document = Document()
    document.add_heading('Secret Santa', 0)
    document.add_paragraph('You are: ' + name)
    document.add_paragraph('Your assigned people to get presents for: ')
    document.add_paragraph(person1_choices[0], style='List Bullet')
    document.add_paragraph(person1_choices[1], style='List Bullet')
    document.add_paragraph(person1_choices[2], style='List Bullet')
    document.add_paragraph('Yay! Merry Xmas! :)')
    document.save(name +"_secret_santa.docx")

'''

'''
dictionary = {'Alex': [], 'Grace': [], 'Charlie': [], 'Karen': [], 'Eric': [], 'Joanne': [], 'Delfina': [], 'Papa and Judy': []}

alex_present_count = 0
grace_present_count = 0
charlie_present_count = 0
karen_present_count = 0
eric_present_count = 0
joanne_present_count = 0
delfina_present_count = 0
papajudy_present_count = 0

for name in names:
	names_copy = names[:]
	names_copy.remove(name)
	choice = random.choice(names_copy)
	i = 0
	if dictionary[]

'''